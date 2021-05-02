#! /usr/bin/env python
# -*- coding: utf-8 -*-

import os
import re
import random
import docx
document = docx.Document()


def run(path, document, can_dirs_regular=[], no_dirs_regular=[], can_files_regular=[], no_files_regular=[], path_start_index=0):
    output_dir = '/'.join(path.split('/')[path_start_index:])
    for file in os.listdir(path):
        file_path = os.path.join(path, file)
        if os.path.isdir(file_path):
            can_output = False
            if len(can_dirs_regular) > 0:
                for can_dir_regular in can_dirs_regular:
                    if re.match(can_dir_regular, file, re.M | re.I):
                        can_output = True
                        break
            else:
                can_output = True
            if not can_output:
                continue

            if len(no_dirs_regular) > 0:
                for no_dir_regular in no_dirs_regular:
                    if re.match(no_dir_regular, file, re.M | re.I):
                        can_output = False
                        break
            if not can_output:
                continue
            run(file_path, document, can_dirs_regular, no_dirs_regular, can_files_regular, no_files_regular, path_start_index)
        else:
            can_output = False
            if len(can_files_regular) > 0:
                for can_file_regular in can_files_regular:
                    if re.match(can_file_regular, file, re.M | re.I):
                        can_output = True
                        break
            else:
                can_output = True
            if not can_output:
                continue

            if len(no_files_regular) > 0:
                for no_file_regular in no_files_regular:
                    if re.match(no_file_regular, file, re.M | re.I):
                        can_output = False
                        break
            if not can_output:
                continue
            with open(file_path, 'r') as f:
                print('输出文件：' + file_path + '\n')
                # f_o.write('文件：' + os.path.join(output_dir, file))
                # f_o.write('\n********************************************************\n')
                # f_o.write(f.read())
                # f_o.write('\n********************************************************\n')
                # f_o.write('\n' * random.randint(1, 3))
                document.add_paragraph('文件：' + os.path.join(output_dir, file))
                table = document.add_table(rows=1, cols=1, style='Table Grid')
                table.cell(0, 0).text = f.read()
                document.add_paragraph('')

    pass


if __name__ == '__main__':
    # path = '/Users/zhengjiaqi/Documents/schoolWork/services/dashuju_front/'
    # # f_o = open('output_front.txt', 'w')
    # can_dirs_regular = []
    # no_dirs_regular = ['\..*', 'dist', 'node_modules', 'pdf', 'test']
    # can_files_regular = ['.*\.js', '.*\.vue', 'Dockerfile', '.*\.html', 'nginx.conf', 'README.md', '.*\.yml']
    # no_files_regular = ['\..*', '.*\.json']
    # path_start_index = 7

    # path = '/Users/zhengjiaqi/Documents/schoolWork/services/jiaowu_system/'
    # # f_o = open('output_service.txt', 'w')
    # can_dirs_regular = []
    # no_dirs_regular = ['\..*', 'cron', 'log', 'migrations', 'tests', '__pycache__']
    # can_files_regular = ['.*\.py', '.*\.yml', '.*\.pem', '.*\.conf', 'Dockerfile', 'README.md', 'requirements.txt']
    # no_files_regular = ['\..*']
    # path_start_index = 6

    path = '/Users/zhengjiaqi/Documents/schoolWork/services/'
    # f_o = open('output_service.txt', 'w')
    can_dirs_regular = []
    no_dirs_regular = ['\..*', 'dashuju_front', 'micro_service_demo', 'test_service', 'docker_test', 'jiaowu_system', 'k8s_yamls', 'tests', '__pycache__']
    can_files_regular = ['.*\.py', '.*\.yml', '.*\.pem', '.*\.conf', 'Dockerfile', 'README.md', 'requirements.txt']
    no_files_regular = ['\..*', 'test.py']
    path_start_index = 6

    run(path, document, can_dirs_regular, no_dirs_regular, can_files_regular, no_files_regular, path_start_index)
    document.save('output_file.docx')
    pass
